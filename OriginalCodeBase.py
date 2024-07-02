# Imports
import tkinter as tk
from tkinter import ttk, simpledialog, messagebox, filedialog
import ttkthemes, os, sys, json, datetime
from datetime import datetime
from docx import Document

# Import JSON Data
JSON_FILE = "BookingData.json"
def Load_JSON():
    try:
        with open(JSON_FILE, "r") as file:
            appData = json.load(file)
            print("SUCCESS: JSON Data Loaded")
            return appData
    except:
        print("ERROR: Unable To Load JSON Data")
        return None
appData = Load_JSON()   

# JSON Data
siteName = appData["SITE_NAME"]
templateDocument = appData["TEMPLATE_DOCUMENT"]
activityRooms = appData["ACTIVITY_ROOMS"]
foodRooms = appData["FOOD_ROOMS"]
partyTypes = appData["PARTY_TYPES"]

# Global Variables
CustomerNameEntry = None,
CustomerEmailEntry = None,
CustomerPhoneEntry = None,
ChildNameEntry = None,
ChildAgeEntry = None,
PartyDateEntry = None,
PartyStartTimeEntry = None,
PartyEndTimeEntry = None,
PartyTypeCheckboxes = [],
PartyRoomCheckboxes = [],
PartyFoodRoomCheckboxes = [],
DateBookedEntry = None,
StaffMemberEntry = None
# Create Application Window

def CreateApp():
    global CustomerNameEntry, CustomerEmailEntry, CustomerPhoneEntry
    global ChildNameEntry, ChildAgeEntry
    global PartyDateEntry, PartyStartTimeEntry, PartyEndTimeEntry, PartyTypeCheckboxes, PartyRoomCheckboxes, PartyFoodRoomCheckboxes
    global DateBookedEntry, StaffMemberEntry
    # Variables
    WINDOW_TITLE = "Booking Confirmation App"
    WINDOW_SIZE = "800x800"
    WINDOW_THEME = "yaru"
    WINDOW_ICON = ""

    # Application Window
    AppWindow = tk.Tk()
    AppWindow.title(WINDOW_TITLE)
    AppWindow.geometry(WINDOW_SIZE)
    AppWindow.resizable(False, False)

    # Application Theme
    AppStyle = ttkthemes.ThemedStyle(AppWindow)
    AppStyle.set_theme(WINDOW_THEME)

    # Window Check
    if AppWindow:
        print("SUCCESS: Application Created")
        print(AppStyle.get_themes())
    else:
        print("ERROR: Unable To Create Application")
        return None
    
    ## ## ## ## ## ## ## ##
    # Create Menu
    ## ## ## ## ## ## ## ##

    # Menu Bar
    MenuBar = tk.Menu(AppWindow)
    AppWindow.config(menu=MenuBar)
    # File Menu
    GeneralMenu = tk.Menu(MenuBar, tearoff=0)
    MenuBar.add_cascade(label="General", menu=GeneralMenu)
    # File Menu - Update Site Name
    GeneralMenu.add_command(label="Update Site Name", command=UpdateSiteName)
    # File Menu - Update Template Document
    GeneralMenu.add_command(label="Update Template Document", command=UpdateTemplateDocument)
    # File Menu - Separator
    GeneralMenu.add_separator()
    # File Menu - Exit
    GeneralMenu.add_command(label="Exit", command=AppWindow.quit)
    # Venue Menu
    VenueMenu = tk.Menu(MenuBar, tearoff=0)
    MenuBar.add_cascade(label="Venue", menu=VenueMenu)
    # File Menu - Update Activity Rooms
    VenueMenu.add_command(label="Update Activity Rooms", command=UpdateActivityRooms)
    # File Menu - Update Food Rooms
    VenueMenu.add_command(label="Update Food Rooms", command=UpdateFoodRooms)
    # File Menu - Update Party Types
    VenueMenu.add_command(label="Update Party Types", command=UpdatePartyTypes)
    
    ## ## ## ## ## ## ## ## 
    # CUSTOMER INFORMATION
    ## ## ## ## ## ## ## ##
    
    # Container
    CustomerInformationContainer = ttk.LabelFrame(AppWindow, text="Customer Information")
    CustomerInformationContainer.pack(padx=0, pady=0, fill="x", side="top", anchor="n")
    # Input Field - Customer Name
    CustomerNameLabel = ttk.Label(CustomerInformationContainer, text="Name:")
    CustomerNameLabel.grid(row=0, column=0, padx=5, pady=5)
    CustomerNameEntry = ttk.Entry(CustomerInformationContainer)
    CustomerNameEntry.grid(row=0, column=1, padx=5, pady=5)
    # Input Field - Customer Email
    CustomerEmailLabel = ttk.Label(CustomerInformationContainer, text="Email:")
    CustomerEmailLabel.grid(row=0, column=2, padx=5, pady=5)
    CustomerEmailEntry = ttk.Entry(CustomerInformationContainer)
    CustomerEmailEntry.grid(row=0, column=3, padx=5, pady=5)
    # Input Field - Customer Phone
    CustomerPhoneLabel = ttk.Label(CustomerInformationContainer, text="Phone:")
    CustomerPhoneLabel.grid(row=0, column=4, padx=5, pady=5)
    CustomerPhoneEntry = ttk.Entry(CustomerInformationContainer)
    CustomerPhoneEntry.grid(row=0, column=5, padx=5, pady=5)
    print("SUCCESS: Customer Information Container Created")

    ## ## ## ## ## ## ## ##
    # Child Information
    ## ## ## ## ## ## ## ##

    # Container
    ChildInformationContainer = ttk.LabelFrame(AppWindow, text="Child Information")
    ChildInformationContainer.pack(padx=0, pady=0, fill="x", after=CustomerInformationContainer)
    # Input Field - Child Name
    ChildNameLabel = ttk.Label(ChildInformationContainer, text="Name:")
    ChildNameLabel.grid(row=0, column=0, padx=5, pady=5)
    ChildNameEntry = ttk.Entry(ChildInformationContainer)
    ChildNameEntry.grid(row=0, column=1, padx=5, pady=5)
    # Input Field - Child Age
    ChildAgeLabel = ttk.Label(ChildInformationContainer, text="Age:")
    ChildAgeLabel.grid(row=0, column=2, padx=5, pady=5)
    ChildAgeEntry = ttk.Entry(ChildInformationContainer)
    ChildAgeEntry.grid(row=0, column=3, padx=5, pady=5)
    print("SUCCESS: Child Information Container Created")

    ## ## ## ## ## ## ## ##
    # Party Information
    ## ## ## ## ## ## ## ##

    # Container
    PartyInformationContainer = ttk.LabelFrame(AppWindow, text="Party Information")
    PartyInformationContainer.pack(padx=0, pady=0, fill="x", after=ChildInformationContainer)
    # Party Time & Date Container
    PartyDateTimeContainer = ttk.LabelFrame(PartyInformationContainer, text="Party Date & Time")
    PartyDateTimeContainer.pack(padx=0, pady=5, fill="x", side="top", anchor="n")
    # Party Date - Label
    PartyDateLabel = ttk.Label(PartyDateTimeContainer, text="Date:")
    PartyDateLabel.grid(row=0, column=0, padx=5, pady=5)
    # Party Date - Entry
    PartyDateEntry = ttk.Entry(PartyDateTimeContainer)
    PartyDateEntry.grid(row=0, column=1, padx=5, pady=5)
    # Party Start Time - Label
    PartyStartTimeLabel = ttk.Label(PartyDateTimeContainer, text="Start Time:")
    PartyStartTimeLabel.grid(row=0, column=2, padx=5, pady=5)
    # Party Start Time - Entry
    PartyStartTimeEntry = ttk.Entry(PartyDateTimeContainer)
    PartyStartTimeEntry.grid(row=0, column=3, padx=5, pady=5)
    # Party End Time - Label
    PartyEndTimeLabel = ttk.Label(PartyDateTimeContainer, text="End Time:")
    PartyEndTimeLabel.grid(row=0, column=4, padx=5, pady=5)
    # Party End Time - Entry
    PartyEndTimeEntry = ttk.Entry(PartyDateTimeContainer)
    PartyEndTimeEntry.grid(row=0, column=5, padx=5, pady=5)
    # Party Type - Container
    PartyTypeInformationContainer = ttk.LabelFrame(PartyInformationContainer, text="Party Type")
    PartyTypeInformationContainer.pack(padx=5, pady=5, fill="x", side="top", anchor="n")
    # Party Type - Checkboxes
    PartyTypeCheckboxes = []
    for partyType in partyTypes:
        partyTypeCheckbox = ttk.Checkbutton(PartyTypeInformationContainer, text=partyType + ": £" + str(partyTypes[partyType]))
        partyTypeCheckbox.pack(padx=5, pady=5, anchor="w", side="top")
        PartyTypeCheckboxes.append(partyTypeCheckbox)
    # Party Room - Container
    PartyRoomInformationContainer = ttk.LabelFrame(PartyInformationContainer, text="Party Room")
    PartyRoomInformationContainer.pack(padx=5, pady=5, fill="x", side="top", anchor="n")
    # Party Room - Checkboxes
    PartyRoomCheckboxes = []
    for partyRoom in activityRooms:
        partyRoomCheckbox = ttk.Checkbutton(PartyRoomInformationContainer, text=partyRoom)
        partyRoomCheckbox.pack(side="left", padx=5, pady=5)
        PartyRoomCheckboxes.append(partyRoomCheckbox)
    # Party Food - Container
    PartyFoodRoomInformationContainer = ttk.LabelFrame(PartyInformationContainer, text="Party Food Room")
    PartyFoodRoomInformationContainer.pack(padx=5, pady=5, fill="x", side="top", anchor="n")
    # Party FoodRoom - Checkboxes
    PartyFoodRoomCheckboxes = []
    for partyFood in foodRooms:
        partyFoodRoomCheckbox = ttk.Checkbutton(PartyFoodRoomInformationContainer, text=partyFood)
        partyFoodRoomCheckbox.pack(side="left", padx=5, pady=5)
        PartyFoodRoomCheckboxes.append(partyFoodRoomCheckbox)
    print("SUCCESS: Party Information Container Created")

    ## ## ## ## ## ## ## ##
    # Admin
    ## ## ## ## ## ## ## ##

    # Container
    AdminContainer = ttk.LabelFrame(AppWindow, text="Admin")
    AdminContainer.pack(padx=0, pady=0, fill="x", after=PartyInformationContainer)
    # Date Booked - Label
    DateBookedLabel = ttk.Label(AdminContainer, text="Date Booked:")
    DateBookedLabel.grid(row=0, column=0, padx=5, pady=5)
    # Date Booked - Entry
    DateBookedEntry = ttk.Entry(AdminContainer)
    DateBookedEntry.grid(row=0, column=1, padx=5, pady=5)
    # Staff Member - Label
    StaffMemberLabel = ttk.Label(AdminContainer, text="Staff Member:")
    StaffMemberLabel.grid(row=0, column=2, padx=5, pady=5)
    # Staff Member - Entry
    StaffMemberEntry = ttk.Entry(AdminContainer)
    StaffMemberEntry.grid(row=0, column=3, padx=5, pady=5)

    ## ## ## ## ## ## ## ##
    # Generate Document Button
    ## ## ## ## ## ## ## ##
    GenerateDocumentButton = ttk.Button(AppWindow, text="Generate Confirmation", command=GenerateDocument)
    GenerateDocumentButton.pack(padx=0, pady=0, fill="x", side="top", anchor="n")

    # Loop Until Closed
    AppWindow.mainloop()

def GenerateDocument():
    global CustomerNameEntry, CustomerEmailEntry, CustomerPhoneEntry
    global ChildNameEntry, ChildAgeEntry
    global PartyDateEntry, PartyStartTimeEntry, PartyEndTimeEntry, PartyTypeCheckboxes, PartyRoomCheckboxes, PartyFoodRoomCheckboxes
    global DateBookedEntry, StaffMemberEntry
    global templateDocument
    if CustomerNameEntry.get() == "":
        messagebox.showerror("Error", "Customer Name: Missing!")
        return
    if CustomerEmailEntry.get() == "":
        messagebox.showerror("Error", "Customer Email: Missing!")
        return
    if CustomerPhoneEntry.get() == "":
        messagebox.showerror("Error", "Customer Contact Number: Missing!")
        return
    if ChildNameEntry.get() == "":
        messagebox.showerror("Error", "Child Name: Missing!")
        return
    if PartyDateEntry.get() == "":
        messagebox.showerror("Error", "Party Date: Missing!")
        return
    if PartyStartTimeEntry.get() == "":
        messagebox.showerror("Error", "Party Start Time: Missing!")
        return
    if PartyEndTimeEntry.get() == "":
        messagebox.showerror("Error", "Party End Time: Missing!")
        return
    if not any(partyType.instate(['selected']) for partyType in PartyTypeCheckboxes):
        messagebox.showerror("Error", "Party Type: Missing!")
        return
    if not any(partyRoom.instate(['selected']) for partyRoom in PartyRoomCheckboxes):
        messagebox.showerror("Error", "Party Room: Missing!")
        return
    if not any(partyFoodRoom.instate(['selected']) for partyFoodRoom in PartyFoodRoomCheckboxes):
        messagebox.showerror("Error", "Party Food Room: Missing!")
        return
    if DateBookedEntry.get() == "":
        messagebox.showerror("Error", "Date Booked: Missing!")
        return
    PARTY_TYPE = []
    for partyType in PartyTypeCheckboxes:
        if partyType.instate(['selected']):
            PARTY_TYPE.append(partyType.cget("text"))
            PARTY_TYPE = PARTY_TYPE[0].split(":")
            PARTY_ACTIVITY = PARTY_TYPE[0].split("-")[0].strip()
            PARTY_COST = PARTY_TYPE[1].replace("£", "")
    PARTY_ROOM = []
    for partyRoom in PartyRoomCheckboxes:
        if partyRoom.instate(['selected']):
            PARTY_ROOM.append(partyRoom.cget("text"))
            PARTY_ROOM = PARTY_ROOM[0]
    PARTY_FOOD_ROOM = []
    for partyFoodRoom in PartyFoodRoomCheckboxes:
        if partyFoodRoom.instate(['selected']):
            PARTY_FOOD_ROOM.append(partyFoodRoom.cget("text"))
            PARTY_FOOD_ROOM = PARTY_FOOD_ROOM[0]

    CUSTOMER_INFORMATION = {
        "CUSTOMER_NAME": CustomerNameEntry.get().split(" ")[0].capitalize() + " " + CustomerNameEntry.get().split(" ")[1].capitalize(),
        "CUSTOMER_EMAIL": CustomerEmailEntry.get(),
        "CUSTOMER_NUMBER": CustomerPhoneEntry.get()
    }
    CHILD_INFORMATION = {
        "CHILD_NAME": ChildNameEntry.get().capitalize(),
        "CHILD_AGE": ChildAgeEntry.get()
    }
    PARTY_INFORMATION = {
        "PARTY_DATE": PartyDateEntry.get(),
        "PARTY_START_TIME": PartyStartTimeEntry.get(),
        "PARTY_END_TIME": PartyEndTimeEntry.get(),
        "PARTY_TYPE": PARTY_ACTIVITY,
        "PARTY_COST": PARTY_COST,
        "PARTY_ROOM": PARTY_ROOM,
        "PARTY_FOOD_ROOM": PARTY_FOOD_ROOM
    }
    ADMIN_INFORMATION = {
        "CUSTOMER_FIRST_NAME": CustomerNameEntry.get().split(" ")[0].capitalize(),
        "DATE_BOOKED": DateBookedEntry.get(),
        "STAFF_MEMBER": StaffMemberEntry.get()
    }
    templateDocument = Document(templateDocument)

    for paragraph in templateDocument.paragraphs:
        for key, value in CUSTOMER_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
        for key, value in PARTY_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        for key, value in CHILD_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        for key, value in ADMIN_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    # Check Tables
    for table in templateDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in CUSTOMER_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                for key, value in PARTY_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                for key, value in CHILD_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                for key, value in ADMIN_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Check Table Columns and Row Cells
    for table in templateDocument.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in PARTY_INFORMATION.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                    for key, value in CHILD_INFORMATION.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                    for key, value in ADMIN_INFORMATION.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))

    # Save the modified document
    saveAsFile = f"{CUSTOMER_INFORMATION['CUSTOMER_NAME']} - {PARTY_ACTIVITY} - Party Confirmation.docx"
    templateDocument.save(saveAsFile)
    print("SUCCESS: Document Saved - ", saveAsFile)
    messagebox.showinfo("Success", f"Document Saved: {saveAsFile}")

    # Clear Fields
    CustomerNameEntry.delete(0, "end")
    CustomerEmailEntry.delete(0, "end")
    CustomerPhoneEntry.delete(0, "end")
    ChildNameEntry.delete(0, "end")
    ChildAgeEntry.delete(0, "end")
    PartyDateEntry.delete(0, "end")
    PartyStartTimeEntry.delete(0, "end")
    PartyEndTimeEntry.delete(0, "end")
    for partyType in PartyTypeCheckboxes:
        partyType.state(['!selected'])
    for partyRoom in PartyRoomCheckboxes:
        partyRoom.state(['!selected'])
    for partyFoodRoom in PartyFoodRoomCheckboxes:
        partyFoodRoom.state(['!selected'])
    DateBookedEntry.delete(0, "end")
    print("SUCCESS: Fields Cleared")

def UpdateTemplateDocument():
    print("SUCCESS: Update Template Document")

def UpdateActivityRooms():
    print("SUCCESS: Update Activity Rooms")

def UpdateFoodRooms():
    print("SUCCESS: Update Food Rooms")

def UpdatePartyTypes():
    print("SUCCESS: Update Party Types")

def UpdateSiteName():
    print("SUCCESS: Update Site Name")

# Main Application
if __name__ == "__main__":
    print("SUCCESS: Application Started")
    # Run Application
    CreateApp()

# End of File