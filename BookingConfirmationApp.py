import sys
import json
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtWidgets import QMessageBox, QDialog, QTableWidget ,QFileDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QCheckBox, QFrame, QGridLayout, QPushButton, QMenuBar, QAction, QMainWindow, QMenu, QAbstractItemView, QTableWidgetItem, QInputDialog
from docx import Document

# Import JSON Data
JSON_FILE = "BookingData.json"

def Load_JSON():
    try:
        with open(JSON_FILE, "r") as file:
            appData = json.load(file)
            print("SUCCESS: JSON Data Loaded")
            return appData
    except:
        print("ERROR: Unable To Load JSON Data")
        return None

appData = Load_JSON()

# JSON Data
siteName = appData["SITE_NAME"]
templateDocument = appData["TEMPLATE_DOCUMENT"]
activityRooms = appData["ACTIVITY_ROOMS"]
foodRooms = appData["FOOD_ROOMS"]
partyTypes = appData["PARTY_TYPES"]

class BookingConfirmationApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle(f"{siteName} - Booking Confirmation Generator")
        self.setGeometry(50, 50, 800, 800)
        
        self.menuBar = self.menuBar()
        self.createMenu()
        self.initUI()


    def initUI(self):
        self.mainLayout = QVBoxLayout()
        self.createCustomerInformation()
        self.createChildInformation()
        self.createPartyInformation()
        self.createAdminInformation()
        self.generateButton = QPushButton("Generate Confirmation")
        self.generateButton.clicked.connect(self.GenerateDocument)
        self.mainLayout.addWidget(self.generateButton)
        centralWidget = QtWidgets.QWidget()
        centralWidget.setLayout(self.mainLayout)
        self.setCentralWidget(centralWidget)

    def createMenu(self):
        generalMenu = self.menuBar.addMenu("General")
        updateSiteName = QAction("Update Site Name", self)
        updateSiteName.triggered.connect(self.UpdateSiteName)
        generalMenu.addAction(updateSiteName)

        updateTemplateDocument = QAction("Update Template Document", self)
        updateTemplateDocument.triggered.connect(self.UpdateTemplateDocument)
        generalMenu.addAction(updateTemplateDocument)

        generalMenu.addSeparator()

        exitApp = QAction("Exit", self)
        exitApp.triggered.connect(self.close)
        generalMenu.addAction(exitApp)

        venueMenu = self.menuBar.addMenu("Venue")
        updateActivityRooms = QAction("Update Activity Rooms", self)
        updateActivityRooms.triggered.connect(self.UpdateActivityRooms)
        venueMenu.addAction(updateActivityRooms)

        updateFoodRooms = QAction("Update Food Rooms", self)
        updateFoodRooms.triggered.connect(self.UpdateFoodRooms)
        venueMenu.addAction(updateFoodRooms)

        updatePartyTypes = QAction("Update Party Types", self)
        updatePartyTypes.triggered.connect(self.UpdatePartyTypes)
        venueMenu.addAction(updatePartyTypes)

    def createCustomerInformation(self):
        customerInfoContainer = QFrame()
        customerInfoLayout = QGridLayout()
        customerInfoContainer.setLayout(customerInfoLayout)

        customerInfoContainer.setFrameShape(QFrame.StyledPanel)
        customerInfoContainer.setFrameShadow(QFrame.Raised)

        customerInfoContainer.setTitle = QLabel("Customer Information")
        customerInfoLayout.addWidget(customerInfoContainer.setTitle, 0, 0, 1, 2)

        self.CustomerNameEntry = QLineEdit()
        self.CustomerEmailEntry = QLineEdit()
        self.CustomerPhoneEntry = QLineEdit()

        customerInfoLayout.addWidget(QLabel("Name:"), 1, 0)
        customerInfoLayout.addWidget(self.CustomerNameEntry, 1, 1)

        customerInfoLayout.addWidget(QLabel("Email:"), 1, 2)
        customerInfoLayout.addWidget(self.CustomerEmailEntry, 1, 3)

        customerInfoLayout.addWidget(QLabel("Phone:"), 1, 4)
        customerInfoLayout.addWidget(self.CustomerPhoneEntry, 1, 5)

        self.mainLayout.addWidget(customerInfoContainer)

    def createChildInformation(self):
        childInfoContainer = QFrame()
        childInfoLayout = QGridLayout()
        childInfoContainer.setLayout(childInfoLayout)

        childInfoContainer.setFrameShape(QFrame.StyledPanel)
        childInfoContainer.setFrameShadow(QFrame.Raised)

        childInfoContainer.setTitle = QLabel("Child Information")
        childInfoLayout.addWidget(childInfoContainer.setTitle, 0, 0, 1, 2)

        self.ChildNameEntry = QLineEdit()
        self.ChildAgeEntry = QLineEdit()

        childInfoLayout.addWidget(QLabel("Name:"), 1, 0)
        childInfoLayout.addWidget(self.ChildNameEntry, 1, 1)

        childInfoLayout.addWidget(QLabel("Age:"), 1, 2)
        childInfoLayout.addWidget(self.ChildAgeEntry, 1, 3)

        self.mainLayout.addWidget(childInfoContainer)

    def createPartyInformation(self):
        partyInfoContainer = QFrame()
        partyInfoLayout = QVBoxLayout()
        partyInfoContainer.setLayout(partyInfoLayout)

        partyInfoContainer.setFrameShape(QFrame.StyledPanel)
        partyInfoContainer.setFrameShadow(QFrame.Raised)

        partyInfoContainer.setTitle = QLabel("Party Information")
        partyInfoLayout.addWidget(partyInfoContainer.setTitle)

        partyDateTimeContainer = QFrame()
        partyDateTimeLayout = QGridLayout()
        partyDateTimeContainer.setLayout(partyDateTimeLayout)

        self.PartyDateEntry = QLineEdit()
        self.PartyStartTimeEntry = QLineEdit()
        self.PartyEndTimeEntry = QLineEdit()

        partyDateTimeLayout.addWidget(QLabel("Date:"), 0, 0)
        partyDateTimeLayout.addWidget(self.PartyDateEntry, 0, 1)

        partyDateTimeLayout.addWidget(QLabel("Start Time:"), 0, 2)
        partyDateTimeLayout.addWidget(self.PartyStartTimeEntry, 0, 3)

        partyDateTimeLayout.addWidget(QLabel("End Time:"), 0, 4)
        partyDateTimeLayout.addWidget(self.PartyEndTimeEntry, 0, 5)

        partyInfoLayout.addWidget(partyDateTimeContainer)

        partyTypeContainer = QFrame()
        partyTypeLayout = QVBoxLayout()
        partyTypeContainer.setLayout(partyTypeLayout)

        partyTypeContainer.setTitle = QLabel("Party Type")
        partyTypeLayout.addWidget(partyTypeContainer.setTitle)

        self.PartyTypeCheckboxes = []
        for partyType, cost in partyTypes.items():
            partyTypeCheckbox = QCheckBox(f"{partyType}: £{cost}")
            partyTypeLayout.addWidget(partyTypeCheckbox)
            self.PartyTypeCheckboxes.append(partyTypeCheckbox)

        partyInfoLayout.addWidget(partyTypeContainer)

        partyRoomContainer = QFrame()
        partyRoomLayout = QHBoxLayout()
        partyRoomContainer.setLayout(partyRoomLayout)

        partyRoomContainer.setTitle = QLabel("Party Room")
        partyRoomLayout.addWidget(partyRoomContainer.setTitle)

        self.PartyRoomCheckboxes = []
        for room in activityRooms:
            partyRoomCheckbox = QCheckBox(room)
            partyRoomLayout.addWidget(partyRoomCheckbox)
            self.PartyRoomCheckboxes.append(partyRoomCheckbox)

        partyInfoLayout.addWidget(partyRoomContainer)

        partyFoodRoomContainer = QFrame()
        partyFoodRoomLayout = QHBoxLayout()
        partyFoodRoomContainer.setLayout(partyFoodRoomLayout)

        partyFoodRoomContainer.setTitle = QLabel("Party Food Room")
        partyFoodRoomLayout.addWidget(partyFoodRoomContainer.setTitle)

        self.PartyFoodRoomCheckboxes = []
        for foodRoom in foodRooms:
            partyFoodRoomCheckbox = QCheckBox(foodRoom)
            partyFoodRoomLayout.addWidget(partyFoodRoomCheckbox)
            self.PartyFoodRoomCheckboxes.append(partyFoodRoomCheckbox)

        partyInfoLayout.addWidget(partyFoodRoomContainer)

        self.mainLayout.addWidget(partyInfoContainer)

    def createAdminInformation(self):
        adminContainer = QFrame()
        adminLayout = QGridLayout()
        adminContainer.setLayout(adminLayout)

        adminContainer.setFrameShape(QFrame.StyledPanel)
        adminContainer.setFrameShadow(QFrame.Raised)

        adminContainer.setTitle = QLabel("Admin")
        adminLayout.addWidget(adminContainer.setTitle, 0, 0, 1, 2)

        self.DateBookedEntry = QLineEdit()
        self.StaffMemberEntry = QLineEdit()

        adminLayout.addWidget(QLabel("Date Booked:"), 1, 0)
        adminLayout.addWidget(self.DateBookedEntry, 1, 1)

        adminLayout.addWidget(QLabel("Staff Member:"), 1, 2)
        adminLayout.addWidget(self.StaffMemberEntry, 1, 3)

        self.mainLayout.addWidget(adminContainer)

    def GenerateDocument(self):
        if self.CustomerNameEntry.text() == "":
            QMessageBox.critical(self, "Error", "Customer Name: Missing!")
            return
        if self.CustomerEmailEntry.text() == "":
            QMessageBox.critical(self, "Error", "Customer Email: Missing!")
            return
        if self.CustomerPhoneEntry.text() == "":
            QMessageBox.critical(self, "Error", "Customer Contact Number: Missing!")
            return
        if self.ChildNameEntry.text() == "":
            QMessageBox.critical(self, "Error", "Child Name: Missing!")
            return
        if self.PartyDateEntry.text() == "":
            QMessageBox.critical(self, "Error", "Party Date: Missing!")
            return
        if self.PartyStartTimeEntry.text() == "":
            QMessageBox.critical(self, "Error", "Party Start Time: Missing!")
            return
        if self.PartyEndTimeEntry.text() == "":
            QMessageBox.critical(self, "Error", "Party End Time: Missing!")
            return
        if not any(cb.isChecked() for cb in self.PartyTypeCheckboxes):
            QMessageBox.critical(self, "Error", "Party Type: Missing!")
            return
        if not any(cb.isChecked() for cb in self.PartyRoomCheckboxes):
            QMessageBox.critical(self, "Error", "Party Room: Missing!")
            return
        if not any(cb.isChecked() for cb in self.PartyFoodRoomCheckboxes):
            QMessageBox.critical(self, "Error", "Party Food Room: Missing!")
            return
        if self.DateBookedEntry.text() == "":
            QMessageBox.critical(self, "Error", "Date Booked: Missing!")
            return

        PARTY_TYPE = [cb.text() for cb in self.PartyTypeCheckboxes if cb.isChecked()][0]
        PARTY_ACTIVITY, PARTY_COST = PARTY_TYPE.split(": £")

        PARTY_ROOM = [cb.text() for cb in self.PartyRoomCheckboxes if cb.isChecked()][0]
        PARTY_FOOD_ROOM = [cb.text() for cb in self.PartyFoodRoomCheckboxes if cb.isChecked()][0]

        CUSTOMER_INFORMATION = {
            "CUSTOMER_NAME": self.CustomerNameEntry.text(),
            "CUSTOMER_EMAIL": self.CustomerEmailEntry.text(),
            "CUSTOMER_NUMBER": self.CustomerPhoneEntry.text()
        }
        CHILD_INFORMATION = {
            "CHILD_NAME": self.ChildNameEntry.text(),
            "CHILD_AGE": self.ChildAgeEntry.text()
        }
        PARTY_INFORMATION = {
            "PARTY_DATE": self.PartyDateEntry.text(),
            "PARTY_START_TIME": self.PartyStartTimeEntry.text(),
            "PARTY_END_TIME": self.PartyEndTimeEntry.text(),
            "PARTY_TYPE": PARTY_ACTIVITY,
            "PARTY_COST": PARTY_COST,
            "PARTY_ROOM": PARTY_ROOM,
            "PARTY_FOOD_ROOM": PARTY_FOOD_ROOM
        }
        ADMIN_INFORMATION = {
            "CUSTOMER_FIRST_NAME": self.CustomerNameEntry.text().split(" ")[0],
            "DATE_BOOKED": self.DateBookedEntry.text(),
            "STAFF_MEMBER": self.StaffMemberEntry.text()
        }

        doc = Document(templateDocument)

        for paragraph in doc.paragraphs:
            for key, value in CUSTOMER_INFORMATION.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
            for key, value in PARTY_INFORMATION.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
            for key, value in CHILD_INFORMATION.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
            for key, value in ADMIN_INFORMATION.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in CUSTOMER_INFORMATION.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                    for key, value in PARTY_INFORMATION.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                    for key, value in CHILD_INFORMATION.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                    for key, value in ADMIN_INFORMATION.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

        saveAsFile = f"{CUSTOMER_INFORMATION['CUSTOMER_NAME']} - {PARTY_ACTIVITY} - Party Confirmation.docx"
        doc.save(saveAsFile)
        print("SUCCESS: Document Saved - ", saveAsFile)
        QMessageBox.information(self, "Success", f"Document Saved: {saveAsFile}")

        # Clear Fields
        self.CustomerNameEntry.clear()
        self.CustomerEmailEntry.clear()
        self.CustomerPhoneEntry.clear()
        self.ChildNameEntry.clear()
        self.ChildAgeEntry.clear()
        self.PartyDateEntry.clear()
        self.PartyStartTimeEntry.clear()
        self.PartyEndTimeEntry.clear()
        for cb in self.PartyTypeCheckboxes:
            cb.setChecked(False)
        for cb in self.PartyRoomCheckboxes:
            cb.setChecked(False)
        for cb in self.PartyFoodRoomCheckboxes:
            cb.setChecked(False)
        self.DateBookedEntry.clear()

    def UpdateTemplateDocument(self):
        print("SUCCESS: Update Template Document")

    def UpdateActivityRooms(self):
        activityRoomsDialog = QDialog(self)
        activityRoomsDialog.setWindowTitle("Update Activity Rooms")
        activityRoomsLayout = QVBoxLayout(activityRoomsDialog)

        activityRoomsTable = QTableWidget()
        activityRoomsTable.setColumnCount(1)
        activityRoomsTable.setHorizontalHeaderLabels(["Activity Room"])
        activityRoomsTable.setEditTriggers(QAbstractItemView.DoubleClicked)
        activityRoomsTable.setRowCount(len(activityRooms))

        for row, activityRoom in enumerate(activityRooms):
            activityRoomItem = QTableWidgetItem(activityRoom)
            activityRoomsTable.setItem(row, 0, activityRoomItem)

        activityRoomsLayout.addWidget(activityRoomsTable)
        activityRoomsTable.resizeColumnsToContents()

        addButton = QPushButton("Add")
        deleteButton = QPushButton("Delete")
        saveButton = QPushButton("Save")
        activityRoomsLayout.addWidget(addButton)
        activityRoomsLayout.addWidget(deleteButton)
        activityRoomsLayout.addWidget(saveButton)

        def addActivityRoom():
            newRow = activityRoomsTable.rowCount()
            activityRoom, ok = QInputDialog.getText(self, "Add Activity Room", "Activity Room:")
            if ok:
                activityRoomsTable.insertRow(newRow)
                activityRoomsTable.setItem(newRow, 0, QTableWidgetItem(activityRoom))

        def deleteActivityRoom():
            selectedRow = activityRoomsTable.currentRow()
            if selectedRow != -1:
                activityRoomsTable.removeRow(selectedRow)

        def saveActivityRooms():
            newActivityRooms = []
            for row in range(activityRoomsTable.rowCount()):
                activityRoom = activityRoomsTable.item(row, 0).text()
                newActivityRooms.append(activityRoom)
            appData["ACTIVITY_ROOMS"] = newActivityRooms
            with open(JSON_FILE, "w") as jsonFile:
                json.dump(appData, jsonFile)
            print("SUCCESS: Activity Rooms Updated")
            QMessageBox.information(self, "Success", "Restart Application For Changes To Take Effect")
            activityRoomsDialog.close()
            self.UpdateAppWindow()

        addButton.clicked.connect(addActivityRoom)
        saveButton.clicked.connect(saveActivityRooms)
        deleteButton.clicked.connect(deleteActivityRoom)

        activityRoomsDialog.exec_()


    def UpdateFoodRooms(self):
        print("SUCCESS: Update Food Rooms")

    def UpdatePartyTypes(self):
        partyTypesDialog = QDialog(self)
        partyTypesDialog.setWindowTitle("Update Party Types")
        partyTypesLayout = QVBoxLayout(partyTypesDialog)
        
        # Create a table to display party types and prices
        partyTypesTable = QTableWidget()
        partyTypesTable.setColumnCount(2)
        partyTypesTable.setHorizontalHeaderLabels(["Party Type", "Party Price"])
        partyTypesTable.setEditTriggers(QAbstractItemView.DoubleClicked)

        # Set Row Length
        partyTypesTable.setRowCount(len(partyTypes))
        
        # Populate the table with existing party types and prices
        for row, (partyType, partyPrice) in enumerate(partyTypes.items()):
            partyTypeItem = QTableWidgetItem(partyType)
            partyPriceItem = QTableWidgetItem(str(partyPrice))
            partyTypesTable.setItem(row, 0, partyTypeItem)
            partyTypesTable.setItem(row, 1, partyPriceItem)
        
        # Add the table to the layout
        partyTypesLayout.addWidget(partyTypesTable)

        partyTypesTable.resizeColumnsToContents()
        
        # Add buttons for adding and saving party types
        addButton = QPushButton("Add")
        deleteButton = QPushButton("Delete")
        saveButton = QPushButton("Save")
        partyTypesLayout.addWidget(addButton)
        partyTypesLayout.addWidget(deleteButton)
        partyTypesLayout.addWidget(saveButton)
        
        # Connect the add button to a function for adding party types
        def addPartyType():
            newRow = partyTypesTable.rowCount()
            partyType, ok = QInputDialog.getText(self, "Add Party Type", "Party Type:")
            if ok:
                partyPrice, ok = QInputDialog.getText(self, "Add Party Price", "Party Price:")
                if ok and partyType and partyPrice:
                    partyTypesTable.insertRow(newRow)
                    partyTypesTable.setItem(newRow, 0, QTableWidgetItem(partyType))
                    partyTypesTable.setItem(newRow, 1, QTableWidgetItem(partyPrice))
        
        def saveParties():
            newPartyTypes = {}
            for row in range(partyTypesTable.rowCount()):
                partyType = partyTypesTable.item(row, 0).text()
                partyPrice = partyTypesTable.item(row, 1).text()
                newPartyTypes[partyType] = partyPrice
            appData["PARTY_TYPES"] = newPartyTypes
            with open(JSON_FILE, "w") as jsonFile:
                json.dump(appData, jsonFile)
            print("SUCCESS: Party Types Updated")
            QMessageBox.information(self, "Success", "Restart Application For Changes To Take Effect")
            partyTypesDialog.close()

        def deletePartyType():
            selectedRow = partyTypesTable.currentRow()
            if selectedRow != -1:
                partyTypesTable.removeRow(selectedRow)
        
        addButton.clicked.connect(addPartyType)
        deleteButton.clicked.connect(deletePartyType)
        saveButton.clicked.connect(saveParties)
        
        # Show the dialog
        partyTypesDialog.exec_()


    def UpdateSiteName(self):
        siteName, ok = QtWidgets.QInputDialog.getText(self, "Update Site Name", "New Site Name:")
        if ok:
            appData["SITE_NAME"] = siteName
            with open(JSON_FILE, "w") as jsonFile:
                json.dump(appData, jsonFile)
            print("SUCCESS: Site Name Updated -", siteName)


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = BookingConfirmationApp()
    window.show()
    sys.exit(app.exec_())