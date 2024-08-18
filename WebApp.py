from flask import Flask, render_template, request, send_file
import json
from docx import Document
import os

app = Flask(__name__)

# Import JSON Data
JSON_FILE = "BookingData.json"

def Load_JSON():
    try:
        with open(JSON_FILE, "r") as file:
            appData = json.load(file)
            print("SUCCESS: JSON Data Loaded")
            return appData
    except Exception as e:
        print(f"ERROR: Unable To Load JSON Data: {e}")
        return None

appData = Load_JSON()

@app.route('/')
def index():
    return render_template('index.html', data=appData)

@app.route('/generate_document', methods=['POST'])
def generate_document():
    # Extract form data
    customer_name = request.form['customer_name']
    customer_email = request.form['customer_email']
    customer_phone = request.form['customer_phone']
    child_name = request.form['child_name']
    child_age = request.form['child_age']
    party_date = request.form['party_date']
    party_start_time = request.form['party_start_time']
    party_end_time = request.form['party_end_time']
    date_booked = request.form['date_booked']
    staff_member = request.form['staff_member']

    party_type = request.form['party_type']
    party_room = request.form['party_room']
    party_food_room = request.form['party_food_room']

    # Split party type and cost
    party_activity, party_cost = party_type.split(": £")

    # Prepare information for document generation
    CUSTOMER_INFORMATION = {
        "CUSTOMER_NAME": customer_name,
        "CUSTOMER_EMAIL": customer_email,
        "CUSTOMER_NUMBER": customer_phone
    }
    CHILD_INFORMATION = {
        "CHILD_NAME": child_name,
        "CHILD_AGE": child_age
    }
    PARTY_INFORMATION = {
        "PARTY_DATE": party_date,
        "PARTY_START_TIME": party_start_time,
        "PARTY_END_TIME": party_end_time,
        "PARTY_TYPE": party_activity,
        "PARTY_COST": party_cost,
        "PARTY_ROOM": party_room,
        "PARTY_FOOD_ROOM": party_food_room,
        "MAX_CHILDREN": appData["MAXIMUM_CHILDREN"][party_activity]
    }
    ADMIN_INFORMATION = {
        "CUSTOMER_FIRST_NAME": customer_name.split(" ")[0],
        "DATE_BOOKED": date_booked,
        "STAFF_MEMBER": staff_member
    }

    # Load the document template
    doc = Document(appData["TEMPLATE_DOCUMENT"])

    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        for key, value in CUSTOMER_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
        for key, value in PARTY_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        for key, value in CHILD_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        for key, value in ADMIN_INFORMATION.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in CUSTOMER_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                for key, value in PARTY_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                for key, value in CHILD_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                for key, value in ADMIN_INFORMATION.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Save the document
    saveAsFile = f"{CUSTOMER_INFORMATION['CUSTOMER_NAME']} - {party_activity} - Party Confirmation.docx"
    doc.save(saveAsFile)
    
    # Send the document as a downloadable file
    return send_file(saveAsFile, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
